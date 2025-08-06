# FastAPI Backend for PDF to DOCX Conversion - Railway Deployment
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
import tempfile
import os
import io
from typing import Optional
import traceback

# Import PDF processing libraries
try:
    from pdf2docx import Converter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

# Initialize FastAPI
app = FastAPI(
    title="PDF to DOCX Converter API",
    description="Convert PDF files to DOCX format via API - Deployed on Railway",
    version="1.0.0"
)

# Configure CORS to allow frontend access
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins for GitHub Pages frontend
    allow_credentials=True,
    allow_methods=["GET", "POST"],
    allow_headers=["*"],
)

# Maximum file size: 100MB
MAX_FILE_SIZE = 100 * 1024 * 1024

@app.get("/")
async def root():
    """Health check and API information"""
    return {
        "service": "PDF to DOCX Converter",
        "status": "healthy",
        "version": "1.0.0",
        "available_methods": {
            "pdf2docx": PDF2DOCX_AVAILABLE,
            "pymupdf": PYMUPDF_AVAILABLE,
            "python_docx": PYTHON_DOCX_AVAILABLE
        },
        "endpoints": {
            "convert": "/convert (POST)",
            "health": "/ (GET)"
        }
    }

@app.post("/convert")
async def convert_pdf_to_docx(file: UploadFile = File(...)):
    """
    Convert uploaded PDF file to DOCX format
    
    Args:
        file: PDF file upload (multipart/form-data)
        
    Returns:
        Downloadable DOCX file
        
    Raises:
        HTTPException: For invalid files, unsupported formats, or conversion errors
    """
    
    # Validate file type
    if not file.filename or not file.filename.lower().endswith('.pdf'):
        raise HTTPException(
            status_code=400, 
            detail="Invalid file type. Only PDF files are supported."
        )
    
    # Read and validate file content
    try:
        content = await file.read()
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Error reading file: {str(e)}"
        )
    
    # Check file size
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Maximum size allowed: {MAX_FILE_SIZE // (1024*1024)}MB"
        )
    
    # Validate PDF content
    if len(content) < 100 or not content.startswith(b'%PDF'):
        raise HTTPException(
            status_code=400,
            detail="Invalid or corrupt PDF file."
        )
    
    # Convert PDF to DOCX
    try:
        docx_content = await convert_pdf_content(content, file.filename)
        
        # Generate output filename
        output_filename = file.filename.rsplit('.', 1)[0] + '.docx'
        
        # Return the DOCX file as a streaming response
        return StreamingResponse(
            io.BytesIO(docx_content),
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f'attachment; filename="{output_filename}"',
                'Content-Length': str(len(docx_content))
            }
        )
        
    except Exception as e:
        # Log the full error for debugging
        error_details = traceback.format_exc()
        print(f"Conversion error: {error_details}")
        
        raise HTTPException(
            status_code=500,
            detail=f"PDF to DOCX conversion failed: {str(e)}"
        )


async def convert_pdf_content(pdf_content: bytes, filename: str) -> bytes:
    """
    Convert PDF content to DOCX using available libraries
    
    Args:
        pdf_content: Raw PDF file content
        filename: Original filename for context
        
    Returns:
        DOCX file content as bytes
        
    Raises:
        Exception: If conversion fails with all available methods
    """
    
    # Create temporary files
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_temp:
        pdf_temp.write(pdf_content)
        pdf_path = pdf_temp.name
    
    try:
        # Method 1: Try pdf2docx (most accurate)
        if PDF2DOCX_AVAILABLE:
            try:
                docx_content = await convert_with_pdf2docx(pdf_path)
                return docx_content
            except Exception as e:
                print(f"pdf2docx conversion failed: {str(e)}")
        
        # Method 2: Try PyMuPDF fallback
        if PYMUPDF_AVAILABLE and PYTHON_DOCX_AVAILABLE:
            try:
                docx_content = await convert_with_pymupdf(pdf_path)
                return docx_content
            except Exception as e:
                print(f"PyMuPDF conversion failed: {str(e)}")
        
        # If no conversion methods are available
        if not PDF2DOCX_AVAILABLE and not PYMUPDF_AVAILABLE:
            raise Exception("No PDF conversion libraries available. Please install pdf2docx or PyMuPDF.")
        
        raise Exception("All conversion methods failed. The PDF might be corrupted or unsupported.")
        
    finally:
        # Clean up temporary PDF file
        if os.path.exists(pdf_path):
            os.unlink(pdf_path)


async def convert_with_pdf2docx(pdf_path: str) -> bytes:
    """
    Convert PDF to DOCX using pdf2docx library
    
    Args:
        pdf_path: Path to temporary PDF file
        
    Returns:
        DOCX content as bytes
    """
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as docx_temp:
        docx_path = docx_temp.name
    
    try:
        # Use pdf2docx converter
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        
        # Read the converted DOCX file
        with open(docx_path, 'rb') as f:
            docx_content = f.read()
        
        return docx_content
        
    finally:
        # Clean up temporary DOCX file
        if os.path.exists(docx_path):
            os.unlink(docx_path)


async def convert_with_pymupdf(pdf_path: str) -> bytes:
    """
    Convert PDF to DOCX using PyMuPDF + python-docx as fallback
    
    Args:
        pdf_path: Path to temporary PDF file
        
    Returns:
        DOCX content as bytes
    """
    
    # Open PDF with PyMuPDF
    pdf_document = fitz.open(pdf_path)
    
    # Create new DOCX document
    doc = Document()
    
    try:
        # Process each page
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            
            # Extract text
            text = page.get_text()
            
            # Add page break if not the first page
            if page_num > 0:
                doc.add_page_break()
            
            # Add page number header
            if text.strip():
                # Add text content
                paragraphs = text.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        paragraph = doc.add_paragraph(para_text.strip())
            else:
                # If no text, add a note about the page
                doc.add_paragraph(f"[Page {page_num + 1} - No extractable text found]")
        
        # Save DOCX to bytes
        with tempfile.NamedTemporaryFile() as docx_temp:
            doc.save(docx_temp.name)
            docx_temp.seek(0)
            docx_content = docx_temp.read()
        
        return docx_content
        
    finally:
        pdf_document.close()


@app.get("/health")
async def health_check():
    """Detailed health check endpoint"""
    return {
        "status": "healthy",
        "libraries": {
            "pdf2docx": PDF2DOCX_AVAILABLE,
            "pymupdf": PYMUPDF_AVAILABLE,
            "python_docx": PYTHON_DOCX_AVAILABLE
        },
        "max_file_size_mb": MAX_FILE_SIZE // (1024 * 1024),
        "supported_formats": {
            "input": ["pdf"],
            "output": ["docx"]
        }
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
