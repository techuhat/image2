# FastAPI Backend for PDF Tools & Image Compression - Render Deployment
from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
import tempfile
import os
import io
from typing import Optional
import traceback
import uuid
from pathlib import Path

# Import PDF processing libraries
try:
    from pdf2docx import Converter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

# Import Image processing libraries
try:
    from PIL import Image, ImageOps
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

# Initialize FastAPI
app = FastAPI(
    title="PDF Tools & Image Compressor API",
    description="Convert PDF to DOCX and compress images",
    version="2.0.0",
    docs_url="/docs",
    redoc_url="/redoc"
)

# CORS Configuration for GitHub Pages
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # For development - in production, specify your domain
    allow_credentials=True,
    allow_methods=["GET", "POST", "OPTIONS"],
    allow_headers=["*"],
)

# Health check endpoint
@app.get("/")
@app.get("/health")
async def health_check():
    """Health check endpoint to verify API is running"""
    return {
        "status": "healthy",
        "message": "PDF Tools & Image Compressor API is running",
        "version": "2.0.0",
        "features": {
            "pdf_to_docx": PDF2DOCX_AVAILABLE or PYMUPDF_AVAILABLE,
            "image_compression": PIL_AVAILABLE
        },
        "libraries": {
            "pdf2docx": PDF2DOCX_AVAILABLE,
            "pymupdf": PYMUPDF_AVAILABLE,
            "python_docx": PYTHON_DOCX_AVAILABLE,
            "pillow": PIL_AVAILABLE
        }
    }

# PDF to DOCX Conversion Endpoint
@app.post("/convert")
async def convert_pdf_to_docx(file: UploadFile = File(...)):
    """
    Convert PDF to DOCX using multiple methods
    Supports files up to 100MB
    """
    # Validate file type
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")
    
    # Check file size (100MB limit)
    file_content = await file.read()
    if len(file_content) > 100 * 1024 * 1024:  # 100MB
        raise HTTPException(status_code=413, detail="File too large. Maximum size is 100MB")
    
    # Reset file pointer
    await file.seek(0)
    
    try:
        # Create temporary files
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_pdf:
            temp_pdf.write(file_content)
            temp_pdf_path = temp_pdf.name
        
        # Generate output filename
        output_filename = f"{file.filename.rsplit('.', 1)[0]}.docx"
        temp_docx_path = tempfile.mktemp(suffix='.docx')
        
        conversion_success = False
        conversion_method = ""
        
        # Method 1: Try pdf2docx (most accurate)
        if PDF2DOCX_AVAILABLE and not conversion_success:
            try:
                converter = Converter(temp_pdf_path)
                converter.convert(temp_docx_path)
                converter.close()
                conversion_success = True
                conversion_method = "pdf2docx"
            except Exception as e:
                print(f"pdf2docx failed: {str(e)}")
        
        # Method 2: Try PyMuPDF + python-docx (fallback)
        if PYMUPDF_AVAILABLE and PYTHON_DOCX_AVAILABLE and not conversion_success:
            try:
                doc = fitz.open(temp_pdf_path)
                word_doc = Document()
                
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text = page.get_text()
                    
                    if text.strip():
                        # Add page content to Word document
                        paragraph = word_doc.add_paragraph(text)
                        
                        # Add page break except for last page
                        if page_num < len(doc) - 1:
                            word_doc.add_page_break()
                    
                    # Try to extract images
                    try:
                        image_list = page.get_images()
                        for img_index, img in enumerate(image_list):
                            xref = img[0]
                            base_image = doc.extract_image(xref)
                            image_bytes = base_image["image"]
                            
                            # Save image temporarily
                            image_temp_path = tempfile.mktemp(suffix='.png')
                            with open(image_temp_path, 'wb') as img_file:
                                img_file.write(image_bytes)
                            
                            # Add image to document
                            word_doc.add_picture(image_temp_path, width=Inches(6))
                            
                            # Clean up temp image
                            os.unlink(image_temp_path)
                    except Exception as img_error:
                        print(f"Image extraction error: {str(img_error)}")
                
                doc.close()
                word_doc.save(temp_docx_path)
                conversion_success = True
                conversion_method = "pymupdf+python-docx"
                
            except Exception as e:
                print(f"PyMuPDF method failed: {str(e)}")
        
        # Method 3: Simple text extraction (last resort)
        if PYMUPDF_AVAILABLE and PYTHON_DOCX_AVAILABLE and not conversion_success:
            try:
                doc = fitz.open(temp_pdf_path)
                word_doc = Document()
                
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text = page.get_text()
                    
                    if text.strip():
                        word_doc.add_paragraph(text)
                        if page_num < len(doc) - 1:
                            word_doc.add_page_break()
                
                doc.close()
                word_doc.save(temp_docx_path)
                conversion_success = True
                conversion_method = "text-only"
                
            except Exception as e:
                print(f"Text-only method failed: {str(e)}")
        
        if not conversion_success:
            raise HTTPException(
                status_code=500, 
                detail="Conversion failed. Please ensure the PDF is not corrupted or password protected."
            )
        
        # Read the converted file
        with open(temp_docx_path, 'rb') as converted_file:
            docx_content = converted_file.read()
        
        # Clean up temporary files
        try:
            os.unlink(temp_pdf_path)
            os.unlink(temp_docx_path)
        except:
            pass
        
        # Return the file
        return StreamingResponse(
            io.BytesIO(docx_content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={output_filename}",
                "X-Conversion-Method": conversion_method
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        # Clean up on error
        try:
            if 'temp_pdf_path' in locals():
                os.unlink(temp_pdf_path)
            if 'temp_docx_path' in locals() and os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
        except:
            pass
        
        raise HTTPException(
            status_code=500, 
            detail=f"Internal server error during conversion: {str(e)}"
        )

# Image Compression Endpoint
@app.post("/compress")
async def compress_image(
    file: UploadFile = File(...),
    quality: int = Form(default=80, description="JPEG quality (1-100)"),
    max_width: Optional[int] = Form(default=None, description="Maximum width in pixels"),
    max_height: Optional[int] = Form(default=None, description="Maximum height in pixels"),
    format: str = Form(default="JPEG", description="Output format (JPEG, PNG, WEBP)")
):
    """
    Compress images with quality and size controls
    Supports: JPEG, PNG, WEBP, BMP, TIFF, GIF
    """
    
    if not PIL_AVAILABLE:
        raise HTTPException(status_code=500, detail="Image processing not available - PIL not installed")
    
    # Validate file type
    allowed_extensions = {'.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff', '.gif'}
    file_ext = Path(file.filename.lower()).suffix
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
        )
    
    # Validate quality
    if not 1 <= quality <= 100:
        raise HTTPException(status_code=400, detail="Quality must be between 1 and 100")
    
    # Validate format
    format = format.upper()
    if format not in ['JPEG', 'PNG', 'WEBP']:
        raise HTTPException(status_code=400, detail="Format must be JPEG, PNG, or WEBP")
    
    # Check file size (50MB limit)
    file_content = await file.read()
    if len(file_content) > 50 * 1024 * 1024:  # 50MB
        raise HTTPException(status_code=413, detail="File too large. Maximum size is 50MB")
    
    try:
        # Open image
        image = Image.open(io.BytesIO(file_content))
        
        # Convert to RGB if necessary (for JPEG output)
        if format == 'JPEG' and image.mode in ('RGBA', 'LA', 'P'):
            # Create white background
            background = Image.new('RGB', image.size, (255, 255, 255))
            if image.mode == 'P':
                image = image.convert('RGBA')
            background.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
            image = background
        
        # Resize if dimensions specified
        if max_width or max_height:
            # Calculate new size maintaining aspect ratio
            width, height = image.size
            
            if max_width and max_height:
                # Fit within both constraints
                ratio = min(max_width / width, max_height / height)
            elif max_width:
                ratio = max_width / width
            else:  # max_height
                ratio = max_height / height
            
            if ratio < 1:  # Only resize if it makes image smaller
                new_width = int(width * ratio)
                new_height = int(height * ratio)
                image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
        
        # Apply optimization
        if format == 'PNG':
            image = ImageOps.optimize(image)
        
        # Compress and save
        output_buffer = io.BytesIO()
        
        save_kwargs = {}
        if format == 'JPEG':
            save_kwargs = {
                'quality': quality,
                'optimize': True,
                'progressive': True
            }
        elif format == 'PNG':
            save_kwargs = {
                'optimize': True,
                'compress_level': 9 - (quality // 11)  # Convert quality to compress_level
            }
        elif format == 'WEBP':
            save_kwargs = {
                'quality': quality,
                'optimize': True,
                'method': 6  # Better compression
            }
        
        image.save(output_buffer, format=format, **save_kwargs)
        compressed_content = output_buffer.getvalue()
        
        # Calculate compression ratio
        original_size = len(file_content)
        compressed_size = len(compressed_content)
        compression_ratio = (1 - compressed_size / original_size) * 100
        
        # Generate output filename
        base_name = Path(file.filename).stem
        extension = '.jpg' if format == 'JPEG' else f'.{format.lower()}'
        output_filename = f"{base_name}_compressed{extension}"
        
        # Determine MIME type
        mime_types = {
            'JPEG': 'image/jpeg',
            'PNG': 'image/png',
            'WEBP': 'image/webp'
        }
        
        return StreamingResponse(
            io.BytesIO(compressed_content),
            media_type=mime_types[format],
            headers={
                "Content-Disposition": f"attachment; filename={output_filename}",
                "X-Original-Size": str(original_size),
                "X-Compressed-Size": str(compressed_size),
                "X-Compression-Ratio": f"{compression_ratio:.1f}%",
                "X-Quality": str(quality),
                "X-Format": format
            }
        )
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Image compression failed: {str(e)}"
        )

# Batch processing endpoint (future enhancement)
@app.post("/batch-compress")
async def batch_compress_images(files: list[UploadFile] = File(...)):
    """
    Batch compress multiple images
    Returns a ZIP file with compressed images
    """
    if not PIL_AVAILABLE:
        raise HTTPException(status_code=500, detail="Image processing not available")
    
    if len(files) > 10:
        raise HTTPException(status_code=400, detail="Maximum 10 files allowed in batch")
    
    return {"message": "Batch processing - Coming soon!", "files_received": len(files)}

# Get API capabilities
@app.get("/capabilities")
async def get_capabilities():
    """Get information about available features and supported formats"""
    return {
        "pdf_conversion": {
            "available": PDF2DOCX_AVAILABLE or PYMUPDF_AVAILABLE,
            "methods": {
                "pdf2docx": PDF2DOCX_AVAILABLE,
                "pymupdf": PYMUPDF_AVAILABLE,
                "python_docx": PYTHON_DOCX_AVAILABLE
            },
            "max_file_size": "100MB",
            "supported_input": ["PDF"],
            "supported_output": ["DOCX"]
        },
        "image_compression": {
            "available": PIL_AVAILABLE,
            "max_file_size": "50MB",
            "supported_input": ["JPEG", "PNG", "WEBP", "BMP", "TIFF", "GIF"],
            "supported_output": ["JPEG", "PNG", "WEBP"],
            "features": ["Quality control", "Resize", "Format conversion", "Optimization"]
        }
    }

# Error handlers
@app.exception_handler(413)
async def payload_too_large_handler(request, exc):
    return JSONResponse(
        status_code=413,
        content={"detail": "File too large"}
    )

@app.exception_handler(422)
async def validation_exception_handler(request, exc):
    return JSONResponse(
        status_code=422,
        content={"detail": "Invalid input parameters"}
    )

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
