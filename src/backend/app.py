#!/usr/bin/env python3
"""
Enhanced Flask Backend for PDF Tools - Production Ready
Version 3.0 - Bug Fixes, Security Enhancements, Performance Optimizations
"""

import os
import tempfile
import subprocess
import hashlib
import logging
import traceback
import time
import signal
import threading
import mimetypes
import platform
from datetime import datetime, timedelta
from pathlib import Path
import zipfile
import shutil
import json
from functools import wraps
from contextlib import contextmanager
import asyncio
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass
from collections import defaultdict, deque
from threading import Lock

# Configuration Management
@dataclass
class Config:
    """Production-ready configuration management"""
    MAX_FILE_SIZE: int = int(os.getenv('MAX_FILE_SIZE', 500 * 1024 * 1024))
    MAX_CONCURRENT_OPERATIONS: int = int(os.getenv('MAX_CONCURRENT_OPS', 5))
    MAX_MEMORY_USAGE: int = int(os.getenv('MAX_MEMORY_USAGE', 80))
    CACHE_MAX_AGE_HOURS: int = int(os.getenv('CACHE_MAX_AGE', 24))
    LOG_LEVEL: str = os.getenv('LOG_LEVEL', 'INFO')
    RATE_LIMIT_PER_MINUTE: int = int(os.getenv('RATE_LIMIT_PER_MINUTE', 20))
    RATE_LIMIT_PER_HOUR: int = int(os.getenv('RATE_LIMIT_PER_HOUR', 100))

config = Config()

# Rate limiting implementation
class RateLimiter:
    def __init__(self):
        self.requests = defaultdict(deque)
        self.lock = Lock()
        self.limits = {
            'default': (config.RATE_LIMIT_PER_HOUR, 3600),  # Per hour
            'pdf_convert': (10, 60),  # 10 PDF conversions per minute  
            'batch_process': (5, 300),  # 5 batch operations per 5 minutes
            'compress': (config.RATE_LIMIT_PER_MINUTE, 60)  # Per minute
        }
    
    def is_allowed(self, client_ip: str, endpoint: str = 'default') -> tuple:
        """Check if request is allowed and return rate limit info"""
        limit_count, limit_window = self.limits.get(endpoint, self.limits['default'])
        now = time.time()
        key = f"{client_ip}:{endpoint}"
        
        with self.lock:
            # Clean old requests
            while (self.requests[key] and 
                   self.requests[key][0] < now - limit_window):
                self.requests[key].popleft()
            
            current_count = len(self.requests[key])
            
            if current_count >= limit_count:
                # Rate limit exceeded
                oldest_request = self.requests[key][0] if self.requests[key] else now
                reset_time = oldest_request + limit_window
                return False, {
                    'limit': limit_count,
                    'remaining': 0,
                    'reset': int(reset_time),
                    'retry_after': int(reset_time - now)
                }
            
            # Allow request
            self.requests[key].append(now)
            return True, {
                'limit': limit_count,
                'remaining': limit_count - current_count - 1,
                'reset': int(now + limit_window),
                'retry_after': 0
            }

# Request monitoring
class RequestMonitor:
    def __init__(self):
        self.requests = deque(maxlen=1000)  # Keep last 1000 requests
        self.lock = Lock()
    
    def log_request(self, endpoint: str, method: str, status_code: int, 
                   duration: float, client_ip: str, user_agent: str = None):
        """Log request details for monitoring"""
        request_data = {
            'timestamp': datetime.now().isoformat(),
            'endpoint': endpoint,
            'method': method,
            'status_code': status_code,
            'duration_ms': round(duration * 1000, 2),
            'client_ip': client_ip,
            'user_agent': user_agent[:100] if user_agent else None
        }
        
        with self.lock:
            self.requests.append(request_data)
    
    def get_stats(self, hours: int = 1) -> dict:
        """Get request statistics for the last N hours"""
        cutoff = datetime.now() - timedelta(hours=hours)
        
        with self.lock:
            recent_requests = [
                req for req in self.requests
                if datetime.fromisoformat(req['timestamp']) > cutoff
            ]
        
        if not recent_requests:
            return {'total_requests': 0, 'period_hours': hours}
        
        # Calculate statistics
        total_requests = len(recent_requests)
        avg_duration = sum(req['duration_ms'] for req in recent_requests) / total_requests
        
        status_codes = defaultdict(int)
        endpoints = defaultdict(int)
        methods = defaultdict(int)
        
        for req in recent_requests:
            status_codes[req['status_code']] += 1
            endpoints[req['endpoint']] += 1
            methods[req['method']] += 1
        
        return {
            'total_requests': total_requests,
            'period_hours': hours,
            'requests_per_hour': round(total_requests / hours, 2),
            'average_duration_ms': round(avg_duration, 2),
            'status_codes': dict(status_codes),
            'endpoints': dict(sorted(endpoints.items(), key=lambda x: x[1], reverse=True)),
            'methods': dict(methods)
        }

# Enhanced logging setup
logging.basicConfig(
    level=getattr(logging, config.LOG_LEVEL),
    format='%(asctime)s - %(levelname)s - %(name)s - [%(funcName)s:%(lineno)d] - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('app.log', mode='a')  # File logging for debugging
    ]
)
logger = logging.getLogger(__name__)

# Security configurations using Config
MAX_FILE_SIZE = config.MAX_FILE_SIZE
MAX_MEMORY_USAGE = config.MAX_MEMORY_USAGE
MAX_CONCURRENT_OPERATIONS = config.MAX_CONCURRENT_OPERATIONS
ALLOWED_MIME_TYPES = {
    'application/pdf': ['.pdf'],
    'image/jpeg': ['.jpg', '.jpeg'],
    'image/png': ['.png'],
    'image/webp': ['.webp'],
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ['.docx']
}

# Enhanced error tracking with rate limiting
error_counts = {}
last_errors = []
MAX_ERROR_HISTORY = 200
operation_semaphore = threading.Semaphore(MAX_CONCURRENT_OPERATIONS)

# Try to import optional dependencies with fallbacks
try:
    import magic  # For better file type detection
    MAGIC_AVAILABLE = True
    logger.info("✅ python-magic loaded successfully")
except ImportError:
    MAGIC_AVAILABLE = False
    logger.warning("❌ python-magic not available, using basic file validation")

try:
    import psutil  # For resource monitoring
    PSUTIL_AVAILABLE = True
    logger.info("✅ psutil loaded successfully")
except ImportError:
    PSUTIL_AVAILABLE = False
    logger.warning("❌ psutil not available, resource monitoring disabled")

class SecurityValidator:
    """Enhanced security validation for file uploads"""
    
    @staticmethod
    def validate_file_signature(file_content, expected_extension):
        """Validate file using magic bytes"""
        if not MAGIC_AVAILABLE:
            return True, "Magic bytes validation not available"
            
        try:
            mime_type = magic.from_buffer(file_content, mime=True)
            if mime_type not in ALLOWED_MIME_TYPES:
                return False, f"Unauthorized MIME type: {mime_type}"
            
            if expected_extension not in ALLOWED_MIME_TYPES[mime_type]:
                return False, f"Extension {expected_extension} doesn't match MIME type {mime_type}"
            
            return True, "Valid file"
        except Exception as e:
            logger.warning(f"Magic bytes validation failed: {e}")
            return True, "Could not validate file signature"
    
    @staticmethod
    def sanitize_filename(filename):
        """Enhanced filename sanitization"""
        import re
        
        # Remove path components
        filename = os.path.basename(filename)
        
        # Replace dangerous characters
        filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        
        # Limit length
        if len(filename) > 100:
            name, ext = os.path.splitext(filename)
            filename = name[:95] + ext
        
        # Prevent hidden files
        if filename.startswith('.'):
            filename = '_' + filename[1:]
        
        return filename

class ResourceMonitor:
    """Monitor system resources and prevent overload"""
    
    @staticmethod
    def check_system_health():
        """Check if system can handle new operations"""
        if not PSUTIL_AVAILABLE:
            return True, "Resource monitoring not available"
            
        try:
            memory = psutil.virtual_memory()
            cpu_percent = psutil.cpu_percent(interval=1)
            disk = psutil.disk_usage('/')
            
            # Check memory usage
            if memory.percent > MAX_MEMORY_USAGE:
                return False, f"High memory usage: {memory.percent}%"
            
            # Check disk space (need at least 1GB free)
            if disk.free < 1024**3:
                return False, f"Low disk space: {disk.free / 1024**3:.1f}GB"
            
            # Check CPU usage
            if cpu_percent > 90:
                return False, f"High CPU usage: {cpu_percent}%"
            
            return True, "System healthy"
            
        except Exception as e:
            logger.warning(f"Resource check failed: {e}")
            return True, "Could not check resources"

def track_error(error_type, error_msg, user_ip="unknown"):
    """Enhanced error tracking with rate limiting"""
    global error_counts, last_errors
    
    current_time = datetime.now()
    error_key = f"{error_type}_{user_ip}"
    
    # Rate limiting: max 10 errors per minute per IP
    recent_errors = [e for e in last_errors 
                    if e.get('ip') == user_ip and 
                    datetime.fromisoformat(e['timestamp']) > current_time - timedelta(minutes=1)]
    
    if len(recent_errors) > 10:
        logger.warning(f"Rate limit exceeded for IP {user_ip}")
        return False
    
    error_counts[error_key] = error_counts.get(error_key, 0) + 1
    error_entry = {
        'timestamp': current_time.isoformat(),
        'type': error_type,
        'message': str(error_msg),
        'ip': user_ip,
        'count': error_counts[error_key]
    }
    
    last_errors.append(error_entry)
    if len(last_errors) > MAX_ERROR_HISTORY:
        last_errors.pop(0)
    
    logger.error(f"Error tracked: {error_type} - {error_msg} (IP: {user_ip})")
    return True

def enhanced_error_handler(func):
    """Production-ready error handling with monitoring"""
    @wraps(func)
    def wrapper(*args, **kwargs):
        from flask import request
        user_ip = request.remote_addr if hasattr(request, 'remote_addr') else 'unknown'
        
        # Check system health before processing
        healthy, health_msg = ResourceMonitor.check_system_health()
        if not healthy:
            from flask import jsonify
            return jsonify({
                "error": "System overloaded",
                "code": "SYSTEM_OVERLOAD",
                "details": health_msg,
                "retry_after": 60
            }), 503
        
        # Acquire semaphore for concurrency control
        if not operation_semaphore.acquire(blocking=False):
            from flask import jsonify
            return jsonify({
                "error": "Server busy",
                "code": "TOO_MANY_REQUESTS",
                "details": f"Maximum {MAX_CONCURRENT_OPERATIONS} concurrent operations allowed",
                "retry_after": 30
            }), 429
        
        try:
            return func(*args, **kwargs)
        except MemoryError as e:
            track_error("MEMORY_ERROR", str(e), user_ip)
            from flask import jsonify
            return jsonify({
                "error": "Insufficient memory",
                "code": "MEMORY_ERROR",
                "suggestion": "Try with a smaller file or reduce quality settings",
                "retry_after": 60
            }), 507
        except subprocess.TimeoutExpired as e:
            track_error("TIMEOUT_ERROR", str(e), user_ip)
            from flask import jsonify
            return jsonify({
                "error": "Operation timed out",
                "code": "TIMEOUT_ERROR", 
                "suggestion": "File may be too large or corrupted",
                "retry_after": 30
            }), 408
        except FileNotFoundError as e:
            track_error("FILE_NOT_FOUND", str(e), user_ip)
            from flask import jsonify
            return jsonify({
                "error": "Required system tool not found",
                "code": "TOOL_MISSING",
                "suggestion": "Server configuration issue",
                "details": str(e)
            }), 500
        except PermissionError as e:
            track_error("PERMISSION_ERROR", str(e), user_ip)
            from flask import jsonify
            return jsonify({
                "error": "Permission denied",
                "code": "PERMISSION_ERROR",
                "suggestion": "Server permission issue"
            }), 403
        except Exception as e:
            track_error("UNKNOWN_ERROR", str(e), user_ip)
            logger.error(f"Unexpected error in {func.__name__}: {e}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            from flask import jsonify
            return jsonify({
                "error": "Internal server error",
                "code": "INTERNAL_ERROR",
                "message": str(e),
                "suggestion": "Please try again or contact support",
                "function": func.__name__
            }), 500
        finally:
            operation_semaphore.release()
    
    return wrapper

def validate_file_size_streaming(file_stream, max_size):
    """Better file size validation with streaming to prevent memory issues"""
    total_size = 0
    chunk_size = 8192
    original_position = file_stream.tell()
    file_stream.seek(0)
    
    try:
        while True:
            chunk = file_stream.read(chunk_size)
            if not chunk:
                break
            total_size += len(chunk)
            if total_size > max_size:
                raise ValueError(f"File too large: {total_size} bytes > {max_size} bytes")
        
        file_stream.seek(0)
        return total_size
    except Exception as e:
        file_stream.seek(original_position)
        raise

def validate_quality_param(quality_str, default=85):
    """Validate and sanitize quality parameter"""
    try:
        quality = int(quality_str)
        if not 1 <= quality <= 100:
            logger.warning(f"Invalid quality {quality}, using default {default}")
            return default
        return quality
    except (ValueError, TypeError):
        logger.warning(f"Invalid quality format '{quality_str}', using default {default}")
        return default

@contextmanager
def temp_file_manager(*file_paths):
    """Context manager for better temp file cleanup"""
    try:
        yield file_paths
    finally:
        for file_path in file_paths:
            try:
                if file_path and Path(file_path).exists():
                    Path(file_path).unlink()
                    logger.debug(f"Cleaned up temp file: {file_path}")
            except Exception as e:
                logger.warning(f"Cleanup failed for {file_path}: {e}")

def validate_file_upload(file, allowed_extensions, max_size_mb=100):
    """Enhanced file validation with security checks and streaming validation"""
    if not file or file.filename == '':
        raise ValueError("No file provided")
    
    # Sanitize filename
    safe_filename = SecurityValidator.sanitize_filename(file.filename)
    if not safe_filename:
        raise ValueError("Invalid filename")
    
    # Check file extension
    file_ext = Path(safe_filename).suffix.lower()
    if file_ext not in allowed_extensions:
        raise ValueError(f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}")
    
    # Use streaming validation for file size
    max_size_bytes = max_size_mb * 1024 * 1024
    try:
        file_size = validate_file_size_streaming(file, max_size_bytes)
    except ValueError as e:
        raise ValueError(f"File validation failed: {str(e)}")
    
    if file_size == 0:
        raise ValueError("File is empty")
    
    # Read first part for magic bytes validation
    file_start = file.read(8192)
    file.seek(0)
    
    # Validate file signature
    is_valid, validation_msg = SecurityValidator.validate_file_signature(file_start, file_ext)
    if not is_valid:
        raise ValueError(f"Security validation failed: {validation_msg}")
    
    return safe_filename

# Flask and web framework imports
from flask import Flask, request, jsonify, send_file, after_this_request, make_response, g
from flask_cors import CORS
from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge

# All PDF and image processing imports with error handling
try:
    from PIL import Image, ImageOps
    PIL_AVAILABLE = True
    logger.info("✅ PIL/Pillow loaded successfully")
except ImportError:
    PIL_AVAILABLE = False
    logger.warning("❌ PIL/Pillow not available")

try:
    from pdf2docx import Converter
    PDF2DOCX_AVAILABLE = True
    logger.info("✅ pdf2docx loaded successfully")
except ImportError:
    PDF2DOCX_AVAILABLE = False
    logger.warning("❌ pdf2docx not available")

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
    logger.info("✅ PyMuPDF loaded successfully")
except ImportError:
    PYMUPDF_AVAILABLE = False
    logger.warning("❌ PyMuPDF not available")

try:
    import pikepdf
    PIKEPDF_AVAILABLE = True
    logger.info("✅ pikepdf loaded successfully")
except ImportError:
    PIKEPDF_AVAILABLE = False
    logger.warning("❌ pikepdf not available")

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
    logger.info("✅ python-docx loaded successfully")
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logger.warning("❌ python-docx not available")

# Initialize Flask app with enhanced configuration
app = Flask(__name__)
app.config.update({
    'MAX_CONTENT_LENGTH': MAX_FILE_SIZE,
    'UPLOAD_FOLDER': 'uploads',
    'SECRET_KEY': os.environ.get('SECRET_KEY') or os.urandom(32).hex(),
    'JSON_SORT_KEYS': False,
    'JSONIFY_PRETTYPRINT_REGULAR': False
})

# Initialize rate limiter and request monitor
rate_limiter = RateLimiter()
request_monitor = RequestMonitor()

# Track application start time for uptime metrics
app.start_time = time.time()

def rate_limit(endpoint: str = 'default'):
    """Rate limiting decorator"""
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            client_ip = request.environ.get('HTTP_X_FORWARDED_FOR', 
                                          request.environ.get('HTTP_X_REAL_IP', 
                                          request.remote_addr))
            
            allowed, rate_info = rate_limiter.is_allowed(client_ip, endpoint)
            
            if not allowed:
                response = jsonify({
                    'error': 'Rate limit exceeded',
                    'message': f'Too many requests for {endpoint}',
                    'retry_after': rate_info['retry_after']
                })
                response.status_code = 429
                response.headers['X-RateLimit-Limit'] = str(rate_info['limit'])
                response.headers['X-RateLimit-Remaining'] = str(rate_info['remaining'])
                response.headers['X-RateLimit-Reset'] = str(rate_info['reset'])
                response.headers['Retry-After'] = str(rate_info['retry_after'])
                return response
            
            # Add rate limit headers to successful responses
            response = make_response(f(*args, **kwargs))
            response.headers['X-RateLimit-Limit'] = str(rate_info['limit'])
            response.headers['X-RateLimit-Remaining'] = str(rate_info['remaining'])
            response.headers['X-RateLimit-Reset'] = str(rate_info['reset'])
            
            return response
        return decorated_function
    return decorator

@app.before_request
def before_request():
    """Set request start time for monitoring"""
    g.start_time = time.time()

@app.after_request  
def after_request(response):
    """Log request details for monitoring"""
    if hasattr(g, 'start_time'):
        duration = time.time() - g.start_time
        client_ip = request.environ.get('HTTP_X_FORWARDED_FOR',
                                      request.environ.get('HTTP_X_REAL_IP', 
                                      request.remote_addr))
        user_agent = request.headers.get('User-Agent')
        
        request_monitor.log_request(
            endpoint=request.endpoint or 'unknown',
            method=request.method,
            status_code=response.status_code,
            duration=duration,
            client_ip=client_ip,
            user_agent=user_agent
        )
    
    return response

# Enhanced CORS configuration
CORS(app, 
     origins=[
         "https://techuhat.github.io",
         "https://image2.vercel.app", 
         "http://localhost:3000",
         "http://127.0.0.1:3000"
     ],
     methods=['GET', 'POST', 'OPTIONS'],
     allow_headers=['Content-Type', 'Authorization'],
     max_age=3600
)

# Create directories with proper permissions
CACHE_DIR = Path("cache")
TEMP_DIR = Path("temp")
UPLOADS_DIR = Path("uploads")

for directory in [CACHE_DIR, TEMP_DIR, UPLOADS_DIR]:
    directory.mkdir(mode=0o755, exist_ok=True)

# Cache management
class CacheManager:
    """Enhanced cache management with cleanup and size limits"""
    
    @staticmethod
    def get_file_hash(file_content):
        """Generate SHA256 hash with salt"""
        salt = os.environ.get('CACHE_SALT', 'default-salt')
        return hashlib.sha256(salt.encode() + file_content).hexdigest()
    
    @staticmethod
    def cleanup_old_cache(max_age_hours=24):
        """Remove old cache files"""
        try:
            current_time = time.time()
            removed_count = 0
            
            for cache_file in CACHE_DIR.glob('*'):
                if cache_file.is_file():
                    file_age = current_time - cache_file.stat().st_mtime
                    if file_age > max_age_hours * 3600:
                        cache_file.unlink()
                        removed_count += 1
            
            # Also clean temp directory
            for temp_file in TEMP_DIR.glob('*'):
                if temp_file.is_file():
                    file_age = current_time - temp_file.stat().st_mtime
                    if file_age > 3600:  # Remove temp files older than 1 hour
                        temp_file.unlink()
                        removed_count += 1
            
            if removed_count > 0:
                logger.info(f"Cleaned up {removed_count} old cache/temp files")
                
        except Exception as e:
            logger.warning(f"Cache cleanup error: {e}")
    
    @staticmethod
    def get_cache_stats():
        """Get cache directory statistics"""
        try:
            cache_files = list(CACHE_DIR.glob('*'))
            temp_files = list(TEMP_DIR.glob('*'))
            
            cache_size = sum(f.stat().st_size for f in cache_files if f.is_file())
            temp_size = sum(f.stat().st_size for f in temp_files if f.is_file())
            
            return {
                'cache_files': len(cache_files),
                'temp_files': len(temp_files),
                'cache_size_mb': round(cache_size / (1024**2), 2),
                'temp_size_mb': round(temp_size / (1024**2), 2),
                'total_size_mb': round((cache_size + temp_size) / (1024**2), 2)
            }
        except Exception as e:
            logger.warning(f"Cache stats error: {e}")
            return {}

# System capability detection with better error handling
def check_system_tools():
    """Comprehensive system tool detection"""
    tools = {}
    
    # Check Ghostscript
    gs_paths = [
        r'C:\Program Files\gs\gs10.05.1\bin\gswin64c.exe',
        r'C:\Program Files\gs\gs9.56.1\bin\gswin64c.exe',
        'gs', 'gswin64c', 'gswin32c'
    ]
    
    for path in gs_paths:
        try:
            result = subprocess.run([path, '--version'], 
                                  capture_output=True, text=True, timeout=5)
            if result.returncode == 0:
                tools['ghostscript'] = {
                    'available': True,
                    'path': path,
                    'version': result.stdout.strip().split('\n')[0] if result.stdout else 'Unknown'
                }
                break
        except (subprocess.SubprocessError, FileNotFoundError, subprocess.TimeoutExpired):
            continue
    
    if 'ghostscript' not in tools:
        tools['ghostscript'] = {'available': False, 'path': None, 'version': None}
    
    # Check LibreOffice
    libreoffice_paths = ['libreoffice', 'soffice']
    for path in libreoffice_paths:
        try:
            result = subprocess.run([path, '--version'], 
                                  capture_output=True, text=True, timeout=10)
            if result.returncode == 0:
                tools['libreoffice'] = {
                    'available': True,
                    'path': path,
                    'version': result.stdout.strip()
                }
                break
        except (subprocess.SubprocessError, FileNotFoundError, subprocess.TimeoutExpired):
            continue
    
    if 'libreoffice' not in tools:
        tools['libreoffice'] = {'available': False, 'path': None, 'version': None}
    
    return tools

# Initialize system tools
SYSTEM_TOOLS = check_system_tools()
logger.info(f"System tools detected: {SYSTEM_TOOLS}")

# Background task for cache cleanup
def background_cleanup():
    """Background task to clean up cache periodically"""
    while True:
        try:
            time.sleep(3600)  # Run every hour
            CacheManager.cleanup_old_cache()
        except Exception as e:
            logger.error(f"Background cleanup error: {e}")

# Start background cleanup thread
cleanup_thread = threading.Thread(target=background_cleanup, daemon=True)
cleanup_thread.start()

# Enhanced health check endpoint
@app.route('/', methods=['GET'])
@app.route('/health', methods=['GET'])
def health_check():
    """Comprehensive health check with enhanced system info"""
    try:
        system_health = ResourceMonitor.check_system_health()
        cache_stats = CacheManager.get_cache_stats()
        
        return jsonify({
            "status": "healthy" if system_health[0] else "degraded",
            "timestamp": datetime.now().isoformat(),
            "version": "3.0.0",
            "message": "Enhanced PDF Toolkit Backend - Production Ready",
            "system_health": {
                "healthy": system_health[0],
                "message": system_health[1],
                "concurrent_operations": MAX_CONCURRENT_OPERATIONS - operation_semaphore._value,
                "max_concurrent": MAX_CONCURRENT_OPERATIONS
            },
            "capabilities": {
                "pdf_compression": SYSTEM_TOOLS['ghostscript']['available'] or PYMUPDF_AVAILABLE,
                "pdf_to_docx": PDF2DOCX_AVAILABLE or (PYMUPDF_AVAILABLE and PYTHON_DOCX_AVAILABLE),
                "image_processing": PIL_AVAILABLE,
                "batch_processing": True,
                "advanced_security": True,
                "resource_monitoring": PSUTIL_AVAILABLE,
                "cache_management": True
            },
            "libraries": {
                "ghostscript": SYSTEM_TOOLS['ghostscript'],
                "libreoffice": SYSTEM_TOOLS['libreoffice'],
                "pdf2docx": PDF2DOCX_AVAILABLE,
                "pymupdf": PYMUPDF_AVAILABLE,
                "pikepdf": PIKEPDF_AVAILABLE,
                "python_docx": PYTHON_DOCX_AVAILABLE,
                "pillow": PIL_AVAILABLE,
                "psutil": PSUTIL_AVAILABLE,
                "magic": MAGIC_AVAILABLE
            },
            "cache_stats": cache_stats,
            "features": {
                "advanced_error_handling": True,
                "security_validation": True,
                "resource_monitoring": PSUTIL_AVAILABLE,
                "concurrent_processing": True,
                "background_cleanup": True,
                "rate_limiting": True
            },
            "limits": {
                "max_file_size": f"{MAX_FILE_SIZE // (1024**2)}MB",
                "max_memory_usage": f"{MAX_MEMORY_USAGE}%",
                "max_concurrent_ops": MAX_CONCURRENT_OPERATIONS,
                "supported_formats": list(ALLOWED_MIME_TYPES.keys())
            }
        })
    except Exception as e:
        logger.error(f"Health check error: {e}")
        return jsonify({
            "status": "error",
            "timestamp": datetime.now().isoformat(),
            "error": str(e)
        }), 500

@app.route('/ready', methods=['GET'])
def readiness_check():
    """Kubernetes/Production readiness probe"""
    try:
        # Check if all required services are available
        if not (PYMUPDF_AVAILABLE or SYSTEM_TOOLS['ghostscript']['available']):
            return jsonify({
                "status": "not ready", 
                "reason": "No PDF processors available",
                "available_libraries": {
                    "pymupdf": PYMUPDF_AVAILABLE,
                    "ghostscript": SYSTEM_TOOLS['ghostscript']['available']
                }
            }), 503
        
        # Check disk space
        if PSUTIL_AVAILABLE:
            try:
                disk = psutil.disk_usage('/')
                free_gb = disk.free / (1024**3)
                if free_gb < 1:  # Less than 1GB
                    return jsonify({
                        "status": "not ready", 
                        "reason": f"Low disk space: {free_gb:.1f}GB available"
                    }), 503
            except Exception as disk_error:
                logger.warning(f"Disk check failed: {disk_error}")
        
        # Check memory usage
        if PSUTIL_AVAILABLE:
            try:
                memory = psutil.virtual_memory()
                if memory.percent > 95:  # Very high memory usage
                    return jsonify({
                        "status": "not ready",
                        "reason": f"High memory usage: {memory.percent}%"
                    }), 503
            except Exception as mem_error:
                logger.warning(f"Memory check failed: {mem_error}")
        
        # Check if directories are writable
        for directory in [CACHE_DIR, TEMP_DIR, UPLOADS_DIR]:
            if not directory.exists():
                return jsonify({
                    "status": "not ready",
                    "reason": f"Required directory missing: {directory}"
                }), 503
        
        return jsonify({
            "status": "ready",
            "timestamp": datetime.now().isoformat(),
            "checks": {
                "pdf_processors": True,
                "disk_space": "ok",
                "memory": "ok",
                "directories": "ok"
            }
        })
        
    except Exception as e:
        logger.error(f"Readiness check error: {e}")
        return jsonify({
            "status": "not ready", 
            "reason": f"Health check failed: {str(e)}"
        }), 503

@app.route('/live', methods=['GET'])
def liveness_check():
    """Kubernetes/Production liveness probe"""
    try:
        # Simple check - if the app is responding, it's alive
        return jsonify({
            "status": "alive",
            "timestamp": datetime.now().isoformat(),
            "uptime_seconds": time.time() - app.start_time if hasattr(app, 'start_time') else 0
        })
    except Exception as e:
        logger.error(f"Liveness check error: {e}")
        return jsonify({
            "status": "error",
            "error": str(e)
        }), 500

# Enhanced endpoint for system statistics
@app.route('/stats', methods=['GET'])
@enhanced_error_handler
def get_system_stats():
    """Enhanced system statistics for monitoring"""
    stats = {
        "timestamp": datetime.now().isoformat(),
        "version": "1.0.0",
        "system": {
            "platform": platform.system(),
            "python_version": platform.python_version(),
            "architecture": platform.architecture()[0]
        },
        "libraries": {
            "pymupdf": PYMUPDF_AVAILABLE,
            "pdf2docx": PDF2DOCX_AVAILABLE,
            "pillow": PIL_AVAILABLE,
            "psutil": PSUTIL_AVAILABLE,
            "magic": MAGIC_AVAILABLE
        },
        "cache": CacheManager.get_cache_stats(),
        "concurrent_operations": concurrent_operations,
        "max_concurrent_operations": MAX_CONCURRENT_OPERATIONS
    }
    
    # Add system resource stats if available
    if PSUTIL_AVAILABLE:
        try:
            # Memory information
            memory = psutil.virtual_memory()
            stats["memory"] = {
                "total_gb": round(memory.total / (1024**3), 2),
                "available_gb": round(memory.available / (1024**3), 2),
                "used_percent": memory.percent,
                "status": "critical" if memory.percent > 90 else "warning" if memory.percent > 75 else "ok"
            }
            
            # CPU information
            stats["cpu"] = {
                "count": psutil.cpu_count(),
                "usage_percent": psutil.cpu_percent(interval=1),
                "load_average": list(psutil.getloadavg()) if hasattr(psutil, 'getloadavg') else None
            }
            
            # Disk information for multiple mount points
            stats["disk"] = {}
            for partition in psutil.disk_partitions():
                try:
                    partition_usage = psutil.disk_usage(partition.mountpoint)
                    stats["disk"][partition.mountpoint] = {
                        "total_gb": round(partition_usage.total / (1024**3), 2),
                        "free_gb": round(partition_usage.free / (1024**3), 2),
                        "used_percent": round((partition_usage.used / partition_usage.total) * 100, 1),
                        "filesystem": partition.fstype
                    }
                except PermissionError:
                    # Skip inaccessible partitions
                    continue
                    
        except Exception as e:
            logger.error(f"Error collecting system stats: {e}")
            stats["system_error"] = str(e)
    
    # Add process-specific stats
    try:
        current_process = psutil.Process() if PSUTIL_AVAILABLE else None
        if current_process:
            with current_process.oneshot():
                stats["process"] = {
                    "pid": current_process.pid,
                    "memory_mb": round(current_process.memory_info().rss / (1024**2), 2),
                    "cpu_percent": current_process.cpu_percent(),
                    "num_threads": current_process.num_threads(),
                    "open_files": len(current_process.open_files()),
                    "connections": len(current_process.connections()),
                    "create_time": datetime.fromtimestamp(current_process.create_time()).isoformat()
                }
    except Exception as e:
        logger.error(f"Error collecting process stats: {e}")
        stats["process_error"] = str(e)
    
    # Add directory stats
    stats["directories"] = {}
    for name, path in [("cache", CACHE_DIR), ("temp", TEMP_DIR), ("uploads", UPLOADS_DIR)]:
        try:
            if path.exists():
                files = list(path.glob("*"))
                total_size = sum(f.stat().st_size for f in files if f.is_file())
                stats["directories"][name] = {
                    "path": str(path),
                    "exists": True,
                    "file_count": len([f for f in files if f.is_file()]),
                    "total_size_mb": round(total_size / (1024**2), 2),
                    "writable": os.access(path, os.W_OK)
                }
            else:
                stats["directories"][name] = {
                    "path": str(path),
                    "exists": False
                }
        except Exception as e:
            stats["directories"][name] = {
                "path": str(path),
                "error": str(e)
            }
    
    return jsonify(stats)

@app.route('/metrics', methods=['GET'])
@enhanced_error_handler  
def get_metrics():
    """Get detailed metrics for monitoring and observability"""
    metrics = {
        "timestamp": datetime.now().isoformat(),
        "system": {
            "uptime_seconds": time.time() - app.start_time if hasattr(app, 'start_time') else 0,
            "version": "1.0.0"
        },
        "requests": {
            "last_hour": request_monitor.get_stats(1),
            "last_24_hours": request_monitor.get_stats(24),
            "rate_limits": {
                endpoint: {
                    "limit": limit[0],
                    "window_seconds": limit[1]
                } for endpoint, limit in rate_limiter.limits.items()
            }
        },
        "operations": {
            "concurrent_operations": concurrent_operations,
            "max_concurrent_operations": MAX_CONCURRENT_OPERATIONS,
            "semaphore_available": operation_semaphore._value
        },
        "cache": CacheManager.get_cache_stats(),
        "libraries": {
            "pymupdf": PYMUPDF_AVAILABLE,
            "pdf2docx": PDF2DOCX_AVAILABLE,
            "pillow": PIL_AVAILABLE,
            "psutil": PSUTIL_AVAILABLE,
            "magic": MAGIC_AVAILABLE
        }
    }
    
    # Add error statistics
    if hasattr(g, 'start_time'):
        recent_errors = defaultdict(int)
        for error in last_errors[-100:]:  # Last 100 errors
            recent_errors[error.get('type', 'unknown')] += 1
        
        metrics["errors"] = {
            "recent_count": len(last_errors[-100:]),
            "by_type": dict(recent_errors),
            "total_tracked": len(last_errors)
        }
    
    return jsonify(metrics)

# Enhanced cache management endpoint
@app.route('/clear-cache', methods=['POST'])
@enhanced_error_handler
def clear_cache():
    """Enhanced cache clearing with options"""
    try:
        clear_type = request.json.get('type', 'all') if request.is_json else 'all'
        
        removed_files = 0
        freed_space = 0
        
        if clear_type in ['all', 'cache']:
            for cache_file in CACHE_DIR.glob('*'):
                if cache_file.is_file():
                    freed_space += cache_file.stat().st_size
                    cache_file.unlink()
                    removed_files += 1
        
        if clear_type in ['all', 'temp']:
            for temp_file in TEMP_DIR.glob('*'):
                if temp_file.is_file():
                    freed_space += temp_file.stat().st_size
                    temp_file.unlink()
                    removed_files += 1
        
        logger.info(f"Cache cleared: {removed_files} files, {freed_space / (1024**2):.1f}MB freed")
        
        return jsonify({
            "message": f"Cache cleared successfully ({clear_type})",
            "files_removed": removed_files,
            "space_freed_mb": round(freed_space / (1024**2), 1),
            "type": clear_type
        })
        
    except Exception as e:
        logger.error(f"Cache clear error: {e}")
        return jsonify({"error": f"Cache clear failed: {str(e)}"}), 500

# Copy all the existing endpoints from app.py with enhanced error handling
@app.route('/pdf-to-docx', methods=['POST'])
@rate_limit('pdf_convert')
@enhanced_error_handler
def pdf_to_docx():
    """
    Convert PDF to DOCX using pdf2docx or PyMuPDF with advanced text and image extraction
    Features: File caching, multiple conversion methods, image preservation
    """
    logger.info("PDF to DOCX conversion request received")
    
    if not (PDF2DOCX_AVAILABLE or PYMUPDF_AVAILABLE):
        return jsonify({
            "error": "PDF to DOCX conversion not available. Install pdf2docx or PyMuPDF"
        }), 500
    
    # Validate file upload
    try:
        validate_file_upload(request.files.get('file'), ['.pdf'], max_size_mb=100)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    
    file = request.files['file']
    force_conversion = request.form.get('force', 'false').lower() == 'true'
    include_images = request.form.get('images', 'true').lower() == 'true'
    
    try:
        # Read file content and generate hash
        file_content = file.read()
        file_hash = CacheManager.get_file_hash(file_content)
        
        # Generate filenames
        original_name = secure_filename(file.filename)
        base_name = Path(original_name).stem
        output_filename = f"{base_name}.docx"
        
        # Check cache
        cache_key = f"{file_hash}_{include_images}_{output_filename}"
        cached_file = CACHE_DIR / cache_key
        
        if cached_file.exists() and not force_conversion:
            logger.info(f"Returning cached DOCX: {output_filename}")
            return send_file(
                cached_file,
                as_attachment=True,
                download_name=output_filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        
        # Create temporary files
        temp_pdf = TEMP_DIR / f"{file_hash}_input.pdf"
        temp_docx = TEMP_DIR / f"{file_hash}_output.docx"
        
        # Write PDF file
        with open(temp_pdf, 'wb') as f:
            f.write(file_content)
        
        conversion_success = False
        conversion_method = ""
        conversion_stats = {"pages": 0, "images": 0, "text_chars": 0}
        
        # Method 1: Try pdf2docx (most accurate with layout preservation)
        if PDF2DOCX_AVAILABLE and not conversion_success:
            try:
                logger.info("Attempting conversion with pdf2docx")
                
                # Configure pdf2docx for better conversion
                converter = Converter(str(temp_pdf))
                
                # Convert with advanced options
                converter.convert(
                    str(temp_docx),
                    start=0,
                    end=None,
                    pages=None,
                    multi_processing=False,  # Disable for stability
                    cpu_count=1
                )
                converter.close()
                
                # Check if conversion was successful
                if temp_docx.exists() and temp_docx.stat().st_size > 0:
                    conversion_success = True
                    conversion_method = "pdf2docx"
                    
                    # Get conversion stats
                    if PYMUPDF_AVAILABLE:
                        doc = fitz.open(str(temp_pdf))
                        conversion_stats["pages"] = len(doc)
                        doc.close()
                    
                    logger.info("Conversion successful with pdf2docx (with layout preservation)")
                else:
                    logger.warning("pdf2docx created empty file")
                    
            except Exception as e:
                logger.warning(f"pdf2docx conversion failed: {e}")
                if temp_docx.exists():
                    temp_docx.unlink()
        
        # Method 2: Try enhanced PyMuPDF with proper DOCX generation
        if PYMUPDF_AVAILABLE and PYTHON_DOCX_AVAILABLE and not conversion_success:
            try:
                logger.info("Attempting conversion with enhanced PyMuPDF")
                
                import io
                from docx.shared import Inches
                
                doc = fitz.open(str(temp_pdf))
                word_doc = Document()
                
                # Add document metadata
                core_props = word_doc.core_properties
                core_props.title = f"Converted from {original_name}"
                core_props.author = "ImagePDF Toolkit"
                core_props.comments = f"Converted using PyMuPDF on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
                
                total_text_chars = 0
                total_images = 0
                
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    
                    # Add page header
                    if page_num > 0:
                        word_doc.add_page_break()
                    
                    # Extract text with better formatting
                    text_dict = page.get_text("dict")
                    
                    # Process text blocks
                    for block in text_dict.get("blocks", []):
                        if "lines" in block:  # Text block
                            block_text = ""
                            for line in block["lines"]:
                                line_text = ""
                                for span in line.get("spans", []):
                                    line_text += span.get("text", "")
                                block_text += line_text + "\n"
                            
                            if block_text.strip():
                                paragraph = word_doc.add_paragraph(block_text.strip())
                                total_text_chars += len(block_text)
                        
                        elif include_images and "image" in block:  # Image block
                            try:
                                # Extract image
                                image_index = block["image"]
                                base_image = doc.extract_image(image_index)
                                image_bytes = base_image["image"]
                                
                                # Add image to document
                                image_stream = io.BytesIO(image_bytes)
                                paragraph = word_doc.add_paragraph()
                                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                                
                                # Calculate image size (max width 6 inches)
                                max_width = Inches(6)
                                run.add_picture(image_stream, width=max_width)
                                
                                total_images += 1
                                logger.info(f"Added image {total_images} from page {page_num + 1}")
                                
                            except Exception as img_error:
                                logger.warning(f"Failed to extract image from page {page_num + 1}: {img_error}")
                                # Add placeholder text for failed image
                                word_doc.add_paragraph(f"[Image {total_images + 1} - extraction failed]")
                
                doc.close()
                
                # Save the document
                word_doc.save(str(temp_docx))
                
                conversion_success = True
                conversion_method = "pymupdf_enhanced"
                conversion_stats = {
                    "pages": len(doc),
                    "images": total_images,
                    "text_chars": total_text_chars
                }
                
                logger.info(f"Enhanced PyMuPDF conversion successful: {conversion_stats}")
                
            except Exception as e:
                logger.warning(f"Enhanced PyMuPDF conversion failed: {e}")
        
        if not conversion_success:
            return jsonify({
                "error": "All conversion methods failed",
                "details": "PDF may be corrupted, password-protected, or contain unsupported content",
                "suggestions": [
                    "Ensure PDF is not password-protected",
                    "Try with a different PDF file",
                    "Check if PDF contains extractable text"
                ]
            }), 500
        
        # Verify output file
        if not temp_docx.exists() or temp_docx.stat().st_size == 0:
            return jsonify({
                "error": "Conversion produced empty file",
                "details": "The PDF might not contain extractable content"
            }), 500
        
        # Copy to cache
        shutil.copy2(temp_docx, cached_file)
        
        # Log successful conversion
        logger.info(f"Successfully converted {original_name} to DOCX using {conversion_method}")
        logger.info(f"Conversion stats: {conversion_stats}")
        
        @after_this_request
        def cleanup(response):
            try:
                if temp_pdf.exists():
                    temp_pdf.unlink()
                if temp_docx.exists():
                    temp_docx.unlink()
            except Exception as e:
                logger.warning(f"Cleanup error: {e}")
            return response
        
        # Return file with conversion stats in headers
        response = send_file(
            cached_file,
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        response.headers['X-Conversion-Method'] = conversion_method
        response.headers['X-Pages-Processed'] = str(conversion_stats['pages'])
        response.headers['X-Images-Extracted'] = str(conversion_stats['images'])
        response.headers['X-Text-Characters'] = str(conversion_stats['text_chars'])
        
        return response
        
    except Exception as e:
        track_error("PDF_TO_DOCX_ERROR", str(e))
        logger.error(f"PDF to DOCX conversion error: {str(e)}")
        return jsonify({
            "error": f"Conversion failed: {str(e)}",
            "code": "CONVERSION_ERROR"
        }), 500

@app.route('/compress-pdf', methods=['POST'])
@rate_limit('compress')
@enhanced_error_handler
def compress_pdf():
    """
    Compress PDF using Ghostscript or PyMuPDF with advanced optimization
    Features: Size optimization, quality control, smart caching, image compression
    """
    logger.info("PDF compression request received")
    
    if not (SYSTEM_TOOLS['ghostscript']['available'] or PYMUPDF_AVAILABLE):
        return jsonify({
            "error": "PDF compression not available. Install Ghostscript or PyMuPDF"
        }), 500
    
    # Validate file upload
    try:
        validate_file_upload(request.files.get('file'), ['.pdf'], max_size_mb=100)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    
    file = request.files['file']
    force_compression = request.form.get('force', 'false').lower() == 'true'
    
    # Quality settings for Ghostscript
    quality_map = {
        'low': '/screen',      # Lowest quality, smallest size
        'medium': '/ebook',    # Medium quality
        'high': '/printer',    # High quality
        'max': '/prepress'     # Maximum quality
    }
    quality = request.form.get('quality', 'medium').lower()
    gs_quality = quality_map.get(quality, '/ebook')
    
    try:
        # Read file content and generate hash
        file_content = file.read()
        file_hash = CacheManager.get_file_hash(file_content)
        original_size = len(file_content)
        
        # Generate filenames
        original_name = secure_filename(file.filename)
        base_name = Path(original_name).stem
        output_filename = f"{base_name}_compressed.pdf"
        
        # Check cache with quality suffix
        cache_key = f"{file_hash}_{quality}_{output_filename}"
        cached_file = CACHE_DIR / cache_key
        
        if cached_file.exists() and not force_compression:
            compressed_size = cached_file.stat().st_size
            compression_ratio = (1 - compressed_size / original_size) * 100
            
            logger.info(f"Returning cached compressed PDF: {output_filename} ({compression_ratio:.1f}% reduction)")
            
            response = send_file(
                cached_file,
                as_attachment=True,
                download_name=output_filename,
                mimetype='application/pdf'
            )
            response.headers['X-Compression-Ratio'] = f"{compression_ratio:.1f}%"
            response.headers['X-Original-Size'] = str(original_size)
            response.headers['X-Compressed-Size'] = str(compressed_size)
            response.headers['X-Cached'] = 'true'
            return response
        
        # Check if file is already small enough
        if original_size < 1024 * 1024 and not force_compression:  # < 1MB
            logger.info("File already small, skipping compression")
            import io
            response = send_file(
                io.BytesIO(file_content),
                as_attachment=True,
                download_name=original_name,
                mimetype='application/pdf'
            )
            response.headers['X-Compression-Ratio'] = "0%"
            response.headers['X-Message'] = "File already optimized"
            return response
        
        # Create temporary files
        temp_input = TEMP_DIR / f"{file_hash}_input.pdf"
        temp_output = TEMP_DIR / f"{file_hash}_compressed.pdf"
        
        # Write input file
        with open(temp_input, 'wb') as f:
            f.write(file_content)
        
        compression_success = False
        compression_method = ""
        compression_stats = {"original_size": original_size, "compressed_size": 0, "method": ""}
        
        # Method 1: Try Ghostscript (preferred for compression)
        if SYSTEM_TOOLS['ghostscript']['available'] and not compression_success:
            try:
                logger.info(f"Attempting compression with Ghostscript (quality: {quality})")
                
                gs_cmd = SYSTEM_TOOLS['ghostscript']['path']
                
                # Enhanced Ghostscript command with better compression
                cmd = [
                    gs_cmd,
                    '-sDEVICE=pdfwrite',
                    '-dCompatibilityLevel=1.4',
                    f'-dPDFSETTINGS={gs_quality}',
                    '-dNOPAUSE',
                    '-dQUIET',
                    '-dBATCH',
                    '-dSAFER',
                    '-dColorImageDownsampleType=/Bicubic',
                    '-dGrayImageDownsampleType=/Bicubic',
                    '-dMonoImageDownsampleType=/Bicubic',
                    '-dOptimize=true',
                    '-dEmbedAllFonts=true',
                    '-dSubsetFonts=true',
                    '-dCompressFonts=true',
                    '-dDetectDuplicateImages=true',
                    f'-sOutputFile={temp_output}',
                    str(temp_input)
                ]
                
                # Add quality-specific settings
                if quality == 'low':
                    cmd.extend([
                        '-dColorImageResolution=72',
                        '-dGrayImageResolution=72',
                        '-dMonoImageResolution=300',
                        '-dDownsampleColorImages=true',
                        '-dDownsampleGrayImages=true'
                    ])
                elif quality == 'medium':
                    cmd.extend([
                        '-dColorImageResolution=150',
                        '-dGrayImageResolution=150',
                        '-dMonoImageResolution=600'
                    ])
                elif quality == 'high':
                    cmd.extend([
                        '-dColorImageResolution=300',
                        '-dGrayImageResolution=300',
                        '-dMonoImageResolution=1200'
                    ])
                
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
                
                if result.returncode == 0 and temp_output.exists() and temp_output.stat().st_size > 0:
                    compression_success = True
                    compression_method = "ghostscript"
                    compression_stats["compressed_size"] = temp_output.stat().st_size
                    compression_stats["method"] = f"Ghostscript ({quality})"
                    logger.info("Compression successful with Ghostscript")
                else:
                    logger.warning(f"Ghostscript compression failed: {result.stderr}")
                    
            except subprocess.TimeoutExpired:
                logger.warning("Ghostscript compression timeout")
                track_error("GHOSTSCRIPT_TIMEOUT", "Ghostscript operation timed out")
            except Exception as e:
                logger.warning(f"Ghostscript compression error: {e}")
                track_error("GHOSTSCRIPT_ERROR", str(e))
        
        # Method 2: Try PyMuPDF (fallback with enhanced compression)
        if PYMUPDF_AVAILABLE and not compression_success:
            try:
                logger.info("Attempting compression with enhanced PyMuPDF")
                
                doc = fitz.open(str(temp_input))
                
                # Quality-based compression parameters
                if quality == 'low':
                    deflate_level = 9
                elif quality == 'medium':
                    deflate_level = 6
                elif quality == 'high':
                    deflate_level = 3
                else:  # max
                    deflate_level = 1
                
                # Advanced compression options
                save_options = {
                    "garbage": 4,  # Remove unused objects
                    "deflate": True,  # Enable deflate compression
                    "deflate_images": True,  # Compress images
                    "deflate_fonts": True,  # Compress fonts
                    "linear": True,  # Optimize for web viewing
                    "clean": True,  # Clean up document structure
                    "pretty": False,  # Don't pretty-print (saves space)
                    "ascii": False,  # Use binary encoding
                    "expand": 0,  # Don't expand content streams
                    "compression": deflate_level
                }
                
                # Save with optimizations
                doc.save(str(temp_output), **save_options)
                doc.close()
                
                if temp_output.exists() and temp_output.stat().st_size > 0:
                    compression_success = True
                    compression_method = "pymupdf_enhanced"
                    compression_stats["compressed_size"] = temp_output.stat().st_size
                    compression_stats["method"] = f"PyMuPDF Enhanced ({quality})"
                    logger.info("Compression successful with enhanced PyMuPDF")
                else:
                    logger.warning("PyMuPDF produced empty file")
                
            except Exception as e:
                logger.warning(f"Enhanced PyMuPDF compression error: {e}")
                track_error("PYMUPDF_ERROR", str(e))
        
        if not compression_success:
            return jsonify({
                "error": "All compression methods failed",
                "details": "PDF may be corrupted, already optimized, or contain unsupported features",
                "suggestions": [
                    "Try with a different quality setting",
                    "Check if PDF is password-protected",
                    "Ensure PDF is not already heavily compressed"
                ]
            }), 500
        
        # Calculate compression ratio
        compressed_size = temp_output.stat().st_size
        compression_ratio = (1 - compressed_size / original_size) * 100
        
        # Only keep compressed version if it's actually smaller
        if compressed_size >= original_size and not force_compression:
            logger.info("Compression didn't reduce file size, returning original")
            response = send_file(
                temp_input,
                as_attachment=True,
                download_name=original_name,
                mimetype='application/pdf'
            )
            response.headers['X-Compression-Ratio'] = "0%"
            response.headers['X-Message'] = "Original file already optimized"
            return response
        
        # Copy to cache
        shutil.copy2(temp_output, cached_file)
        
        # Update compression stats
        compression_stats["compression_ratio"] = compression_ratio
        
        # Log compression
        logger.info(f"Successfully compressed {original_name} using {compression_method}")
        logger.info(f"Compression stats: {compression_stats}")
        
        @after_this_request
        def cleanup(response):
            try:
                if temp_input.exists():
                    temp_input.unlink()
                if temp_output.exists():
                    temp_output.unlink()
            except Exception as e:
                logger.warning(f"Cleanup error: {e}")
            return response
        
        response = send_file(
            cached_file,
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/pdf'
        )
        response.headers['X-Compression-Ratio'] = f"{compression_ratio:.1f}%"
        response.headers['X-Original-Size'] = str(original_size)
        response.headers['X-Compressed-Size'] = str(compressed_size)
        response.headers['X-Method'] = compression_method
        response.headers['X-Quality'] = quality
        
        return response
        
    except Exception as e:
        track_error("PDF_COMPRESSION_ERROR", str(e))
        logger.error(f"PDF compression error: {str(e)}")
        return jsonify({
            "error": f"Compression failed: {str(e)}",
            "code": "COMPRESSION_ERROR"
        }), 500

@app.route('/compress-image', methods=['POST'])
@rate_limit('compress')
@enhanced_error_handler
def compress_image():
    """
    Compress image files using Pillow
    Features: Quality control, format optimization, smart caching
    """
    logger.info("Image compression request received")
    
    if not PIL_AVAILABLE:
        return jsonify({
            "error": "Image compression not available. Install Pillow"
        }), 500
    
    # Validate file upload
    try:
        validate_file_upload(request.files.get('file'), ['.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff'], max_size_mb=50)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    
    file = request.files['file']
    quality = int(request.form.get('quality', 85))  # Default quality 85%
    force_compression = request.form.get('force', 'false').lower() == 'true'
    
    try:
        import io
        
        # Read file content and generate hash
        file_content = file.read()
        file_hash = CacheManager.get_file_hash(file_content)
        original_size = len(file_content)
        
        # Generate filenames
        original_name = secure_filename(file.filename)
        base_name = Path(original_name).stem
        file_ext = Path(original_name).suffix.lower()
        output_filename = f"{base_name}_compressed{file_ext}"
        
        # Check cache with quality suffix
        cached_file = CACHE_DIR / f"{file_hash}_q{quality}_{output_filename}"
        
        if cached_file.exists() and not force_compression:
            compressed_size = cached_file.stat().st_size
            compression_ratio = (1 - compressed_size / original_size) * 100
            
            logger.info(f"Returning cached compressed image: {output_filename} ({compression_ratio:.1f}% reduction)")
            
            response = send_file(
                cached_file,
                as_attachment=True,
                download_name=output_filename,
                mimetype=f'image/{file_ext[1:]}'
            )
            response.headers['X-Compression-Ratio'] = f"{compression_ratio:.1f}%"
            response.headers['X-Original-Size'] = str(original_size)
            response.headers['X-Compressed-Size'] = str(compressed_size)
            return response
        
        # Open and process image
        img = Image.open(io.BytesIO(file_content))
        
        # Convert RGBA to RGB if saving as JPEG
        if file_ext.lower() in ['.jpg', '.jpeg'] and img.mode in ('RGBA', 'LA', 'P'):
            # Create white background for transparency
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
            img = background
        
        # Compress and save
        output_buffer = io.BytesIO()
        
        save_kwargs = {
            'format': img.format if img.format else 'JPEG',
            'quality': quality,
            'optimize': True
        }
        
        # Format specific optimizations
        if file_ext.lower() in ['.jpg', '.jpeg']:
            save_kwargs['progressive'] = True
        elif file_ext.lower() == '.png':
            save_kwargs['compress_level'] = 9
            del save_kwargs['quality']  # PNG doesn't use quality
        elif file_ext.lower() == '.webp':
            save_kwargs['method'] = 6  # Better compression
        
        img.save(output_buffer, **save_kwargs)
        compressed_data = output_buffer.getvalue()
        compressed_size = len(compressed_data)
        
        # Calculate compression ratio
        compression_ratio = (1 - compressed_size / original_size) * 100
        
        # Only keep compressed version if it's actually smaller
        if compressed_size >= original_size and not force_compression:
            logger.info("Compression didn't reduce file size, returning original")
            response = send_file(
                io.BytesIO(file_content),
                as_attachment=True,
                download_name=original_name,
                mimetype=f'image/{file_ext[1:]}'
            )
            response.headers['X-Compression-Ratio'] = "0%"
            response.headers['X-Message'] = "Original file already optimized"
            return response
        
        # Save to cache
        with open(cached_file, 'wb') as f:
            f.write(compressed_data)
        
        # Log compression
        logger.info(f"Successfully compressed {original_name} "
                   f"({compression_ratio:.1f}% reduction: {original_size} -> {compressed_size} bytes)")
        
        response = send_file(
            io.BytesIO(compressed_data),
            as_attachment=True,
            download_name=output_filename,
            mimetype=f'image/{file_ext[1:]}'
        )
        response.headers['X-Compression-Ratio'] = f"{compression_ratio:.1f}%"
        response.headers['X-Original-Size'] = str(original_size)
        response.headers['X-Compressed-Size'] = str(compressed_size)
        
        return response
        
    except Exception as e:
        logger.error(f"Image compression error: {str(e)}")
        return jsonify({"error": f"Compression failed: {str(e)}"}), 500

@app.route('/batch-process', methods=['POST'])
@rate_limit('batch_process')
@enhanced_error_handler
def batch_process():
    """
    Batch process multiple files
    Supports mixed operations and returns ZIP archive
    """
    logger.info("Batch processing request received")
    
    if 'files' not in request.files:
        return jsonify({"error": "No files uploaded"}), 400
    
    files = request.files.getlist('files')
    operation = request.form.get('operation', 'auto')  # auto, compress, convert
    
    if len(files) > 10:
        return jsonify({"error": "Maximum 10 files allowed in batch"}), 400
    
    if not files or all(f.filename == '' for f in files):
        return jsonify({"error": "No files selected"}), 400
    
    try:
        # Create temporary directory for batch processing
        batch_id = hashlib.sha256(str(datetime.now()).encode()).hexdigest()[:8]
        batch_dir = TEMP_DIR / f"batch_{batch_id}"
        batch_dir.mkdir(exist_ok=True)
        
        processed_files = []
        errors = []
        
        for file in files:
            if file.filename == '':
                continue
                
            try:
                original_name = secure_filename(file.filename)
                file_ext = Path(original_name).suffix.lower()
                
                # Determine operation based on file type and user preference
                if operation == 'auto':
                    if file_ext == '.pdf':
                        # Auto: compress PDF
                        current_op = 'compress'
                    elif file_ext in ['.jpg', '.jpeg', '.png', '.webp']:
                        # Auto: compress image
                        current_op = 'compress_image'
                    else:
                        errors.append(f"{original_name}: Unsupported file type")
                        continue
                else:
                    current_op = operation
                
                # Process file based on operation
                if current_op == 'compress' and file_ext == '.pdf':
                    # Simple PDF compression for batch
                    file_content = file.read()
                    temp_file = batch_dir / f"compressed_{original_name}"
                    
                    if PYMUPDF_AVAILABLE:
                        temp_input = batch_dir / f"input_{original_name}"
                        with open(temp_input, 'wb') as f:
                            f.write(file_content)
                        
                        doc = fitz.open(str(temp_input))
                        doc.save(str(temp_file), garbage=4, deflate=True)
                        doc.close()
                        
                        processed_files.append(str(temp_file))
                        temp_input.unlink()
                    else:
                        errors.append(f"{original_name}: PDF compression not available")
                
                elif current_op == 'compress_image' and file_ext in ['.jpg', '.jpeg', '.png', '.webp']:
                    # Simple image compression for batch
                    if PIL_AVAILABLE:
                        file_content = file.read()
                        import io
                        
                        img = Image.open(io.BytesIO(file_content))
                        temp_file = batch_dir / f"compressed_{original_name}"
                        
                        # Convert RGBA to RGB if saving as JPEG
                        if file_ext.lower() in ['.jpg', '.jpeg'] and img.mode in ('RGBA', 'LA', 'P'):
                            background = Image.new('RGB', img.size, (255, 255, 255))
                            if img.mode == 'P':
                                img = img.convert('RGBA')
                            background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                            img = background
                        
                        # Save with compression
                        save_kwargs = {'quality': 85, 'optimize': True}
                        if file_ext.lower() in ['.jpg', '.jpeg']:
                            save_kwargs['progressive'] = True
                        elif file_ext.lower() == '.png':
                            save_kwargs = {'optimize': True, 'compress_level': 9}
                        
                        img.save(str(temp_file), **save_kwargs)
                        processed_files.append(str(temp_file))
                    else:
                        errors.append(f"{original_name}: Image compression not available")
                
                else:
                    errors.append(f"{original_name}: Operation not supported for this file type")
                    
            except Exception as e:
                errors.append(f"{file.filename}: {str(e)}")
        
        if not processed_files:
            return jsonify({
                "error": "No files were successfully processed",
                "errors": errors
            }), 400
        
        # Create ZIP archive
        zip_filename = f"batch_processed_{batch_id}.zip"
        zip_path = batch_dir / zip_filename
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file_path in processed_files:
                file_path_obj = Path(file_path)
                zipf.write(file_path, file_path_obj.name)
        
        logger.info(f"Batch processing completed: {len(processed_files)} files processed, {len(errors)} errors")
        
        @after_this_request
        def cleanup(response):
            try:
                shutil.rmtree(batch_dir)
            except Exception as e:
                logger.warning(f"Batch cleanup error: {e}")
            return response
        
        response = send_file(
            zip_path,
            as_attachment=True,
            download_name=zip_filename,
            mimetype='application/zip'
        )
        response.headers['X-Processed-Files'] = str(len(processed_files))
        response.headers['X-Errors'] = str(len(errors))
        
        return response
        
    except Exception as e:
        logger.error(f"Batch processing error: {str(e)}")
        return jsonify({"error": f"Batch processing failed: {str(e)}"}), 500

@app.route('/pdf-merge', methods=['POST'])
@enhanced_error_handler
def pdf_merge():
    """
    Merge multiple PDF files into a single PDF
    Features: Multiple PDF merging, order preservation
    """
    logger.info("PDF merge request received")
    
    if not PYMUPDF_AVAILABLE:
        return jsonify({
            "error": "PDF merging not available. Install PyMuPDF"
        }), 500
    
    if 'files' not in request.files:
        return jsonify({"error": "No files uploaded"}), 400
    
    files = request.files.getlist('files')
    
    if len(files) < 2:
        return jsonify({"error": "At least 2 PDF files are required for merging"}), 400
    
    if len(files) > 20:
        return jsonify({"error": "Maximum 20 files allowed for merging"}), 400
    
    try:
        # Create merged PDF
        merged_doc = fitz.open()
        
        for file in files:
            if file.filename == '':
                continue
                
            if not file.filename.lower().endswith('.pdf'):
                continue
            
            # Read file content
            file_content = file.read()
            temp_pdf = TEMP_DIR / f"temp_{hash(file_content)}.pdf"
            
            with open(temp_pdf, 'wb') as f:
                f.write(file_content)
            
            # Open and merge
            doc = fitz.open(str(temp_pdf))
            merged_doc.insert_pdf(doc)
            doc.close()
            
            # Cleanup temp file
            if temp_pdf.exists():
                temp_pdf.unlink()
        
        # Save merged PDF
        output_filename = f"merged_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        output_path = TEMP_DIR / output_filename
        
        merged_doc.save(str(output_path))
        merged_doc.close()
        
        logger.info(f"Successfully merged {len(files)} PDF files into {output_filename}")
        
        @after_this_request
        def cleanup(response):
            try:
                if output_path.exists():
                    output_path.unlink()
            except Exception as e:
                logger.warning(f"Cleanup error: {e}")
            return response
        
        return send_file(
            output_path,
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/pdf'
        )
        
    except Exception as e:
        logger.error(f"PDF merge error: {str(e)}")
        return jsonify({"error": f"PDF merging failed: {str(e)}"}), 500

@app.route('/pdf-split', methods=['POST'])
@enhanced_error_handler
def pdf_split():
    """
    Split PDF into multiple files by pages
    Features: Page range splitting, individual page extraction
    """
    logger.info("PDF split request received")
    
    if not PYMUPDF_AVAILABLE:
        return jsonify({
            "error": "PDF splitting not available. Install PyMuPDF"
        }), 500
    
    # Validate file upload
    try:
        validate_file_upload(request.files.get('file'), ['.pdf'], max_size_mb=100)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    
    file = request.files['file']
    split_type = request.form.get('type', 'individual')  # individual, range, custom
    page_ranges = request.form.get('ranges', '')  # e.g., "1-3,5,7-9"
    pages_per_split = request.form.get('pages_per_split', '1')  # for bulk splitting
    
    try:
        # Read file content
        file_content = file.read()
        file_hash = CacheManager.get_file_hash(file_content)
        
        temp_pdf = TEMP_DIR / f"{file_hash}_input.pdf"
        
        with open(temp_pdf, 'wb') as f:
            f.write(file_content)
        
        doc = fitz.open(str(temp_pdf))
        total_pages = len(doc)
        
        if total_pages == 0:
            doc.close()
            return jsonify({
                "error": "PDF contains no pages",
                "details": "The uploaded PDF file appears to be empty"
            }), 400
        
        logger.info(f"Processing PDF with {total_pages} pages, split_type: {split_type}")
        
        # Parse page ranges based on split type
        ranges = []
        
        if split_type == 'individual':
            # Split into individual pages
            ranges = [(i, i) for i in range(total_pages)]
            
        elif split_type == 'bulk' and pages_per_split.isdigit():
            # Split into chunks of specified size
            pages_per_chunk = int(pages_per_split)
            if pages_per_chunk < 1:
                pages_per_chunk = 1
            
            for i in range(0, total_pages, pages_per_chunk):
                end_page = min(i + pages_per_chunk - 1, total_pages - 1)
                ranges.append((i, end_page))
                
        elif split_type == 'range' and page_ranges:
            # Parse custom page ranges
            try:
                for range_str in page_ranges.split(','):
                    range_str = range_str.strip()
                    if '-' in range_str:
                        start_str, end_str = range_str.split('-', 1)
                        start = int(start_str.strip()) - 1  # Convert to 0-based
                        end = int(end_str.strip()) - 1
                        
                        # Validate range
                        if start < 0 or end >= total_pages or start > end:
                            doc.close()
                            return jsonify({
                                "error": f"Invalid page range: {range_str}",
                                "details": f"Valid page numbers are 1-{total_pages}"
                            }), 400
                        
                        ranges.append((start, end))
                    else:
                        page_num = int(range_str) - 1
                        if page_num < 0 or page_num >= total_pages:
                            doc.close()
                            return jsonify({
                                "error": f"Invalid page number: {range_str}",
                                "details": f"Valid page numbers are 1-{total_pages}"
                            }), 400
                        ranges.append((page_num, page_num))
                        
            except ValueError as ve:
                doc.close()
                return jsonify({
                    "error": f"Invalid page range format: {page_ranges}",
                    "details": "Use format like '1-3,5,7-9' (page numbers start from 1)"
                }), 400
        else:
            doc.close()
            return jsonify({
                "error": "Invalid split configuration",
                "details": "Specify valid split_type and parameters"
            }), 400
        
        if not ranges:
            doc.close()
            return jsonify({
                "error": "No valid page ranges specified",
                "details": "Check your page range configuration"
            }), 400
        
        # Create ZIP for multiple files
        original_name = secure_filename(file.filename)
        base_name = Path(original_name).stem
        zip_filename = f"{base_name}_split_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        zip_path = TEMP_DIR / zip_filename
        
        split_files = []
        
        # Process each range
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED, compresslevel=6) as zipf:
            for i, (start, end) in enumerate(ranges):
                try:
                    # Create new PDF with selected pages
                    new_doc = fitz.open()
                    new_doc.insert_pdf(doc, from_page=start, to_page=end)
                    
                    # Generate descriptive filename
                    if start == end:
                        split_filename = f"{base_name}_page_{start+1}.pdf"
                        page_info = f"page {start+1}"
                    else:
                        split_filename = f"{base_name}_pages_{start+1}-{end+1}.pdf"
                        page_info = f"pages {start+1}-{end+1}"
                    
                    # Save to temporary file
                    split_path = TEMP_DIR / split_filename
                    
                    # Save with optimization
                    new_doc.save(
                        str(split_path),
                        garbage=4,  # Remove unused objects
                        deflate=True,  # Compress
                        linear=True  # Optimize for web
                    )
                    new_doc.close()
                    
                    # Verify file was created successfully
                    if split_path.exists() and split_path.stat().st_size > 0:
                        # Add to ZIP
                        zipf.write(split_path, split_filename)
                        split_files.append({
                            "filename": split_filename,
                            "pages": page_info,
                            "size": split_path.stat().st_size
                        })
                        logger.info(f"Created split file: {split_filename} ({page_info})")
                        
                        # Cleanup temp split file
                        split_path.unlink()
                    else:
                        logger.warning(f"Failed to create split file for {page_info}")
                        
                except Exception as split_error:
                    logger.error(f"Error creating split for range {start+1}-{end+1}: {split_error}")
        
        doc.close()
        
        # Check if any splits were successful
        if not split_files:
            return jsonify({
                "error": "Failed to create any split files",
                "details": "All page ranges failed to process"
            }), 500
        
        # Log successful split
        logger.info(f"Successfully split PDF: {len(split_files)} files created")
        
        @after_this_request
        def cleanup(response):
            try:
                if temp_pdf.exists():
                    temp_pdf.unlink()
                if zip_path.exists():
                    zip_path.unlink()
            except Exception as e:
                logger.warning(f"Cleanup error: {e}")
            return response
        
        # Prepare response with metadata
        response = send_file(
            zip_path,
            as_attachment=True,
            download_name=zip_filename,
            mimetype='application/zip'
        )
        
        response.headers['X-Total-Pages'] = str(total_pages)
        response.headers['X-Split-Count'] = str(len(split_files))
        response.headers['X-Split-Type'] = split_type
        
        return response
        
    except Exception as e:
        track_error("PDF_SPLIT_ERROR", str(e))
        logger.error(f"PDF split error: {str(e)}")
        return jsonify({
            "error": f"PDF splitting failed: {str(e)}",
            "code": "SPLIT_ERROR"
        }), 500

# Enhanced error handlers
@app.errorhandler(413)
def too_large(e):
    return jsonify({
        "error": "File too large",
        "code": "FILE_TOO_LARGE",
        "max_size": f"{MAX_FILE_SIZE // (1024**2)}MB",
        "suggestion": "Please use a smaller file"
    }), 413

@app.errorhandler(429)
def rate_limited(e):
    return jsonify({
        "error": "Too many requests",
        "code": "RATE_LIMITED",
        "suggestion": "Please wait before making another request",
        "retry_after": 60
    }), 429

@app.errorhandler(404)
def not_found(e):
    return jsonify({
        "error": "Endpoint not found",
        "code": "NOT_FOUND",
        "available_endpoints": [
            "/health", "/stats", "/pdf-to-docx", "/compress-pdf",
            "/compress-image", "/batch-process", "/pdf-merge", 
            "/pdf-split", "/clear-cache"
        ]
    }), 404

def run_startup_checks():
    """Run comprehensive startup checks"""
    logger.info("🚀 Starting production backend...")
    logger.info(f"📊 Max file size: {MAX_FILE_SIZE / (1024*1024):.1f} MB")
    logger.info(f"⚡ Max concurrent operations: {MAX_CONCURRENT_OPERATIONS}")
    logger.info(f"💾 Cache directory: {CACHE_DIR}")
    logger.info(f"� Temp directory: {TEMP_DIR}")
    
    # Check library availability
    libs = {
        "PyMuPDF": PYMUPDF_AVAILABLE,
        "pdf2docx": PDF2DOCX_AVAILABLE, 
        "Pillow": PIL_AVAILABLE,
        "psutil": PSUTIL_AVAILABLE,
        "python-magic": MAGIC_AVAILABLE
    }
    
    for lib, available in libs.items():
        status = "✅" if available else "❌"
        logger.info(f"{status} {lib}: {'Available' if available else 'Not available'}")
    
    # Verify critical libraries
    if not (PYMUPDF_AVAILABLE or PDF2DOCX_AVAILABLE):
        logger.error("❌ CRITICAL: No PDF processing libraries available!")
        return False
    
    if not PIL_AVAILABLE:
        logger.warning("⚠️  WARNING: Pillow not available, image processing limited")
    
    # Test directory permissions
    for directory in [CACHE_DIR, TEMP_DIR, UPLOADS_DIR]:
        try:
            test_file = directory / f"test_{int(time.time())}.tmp"
            test_file.write_text("test")
            test_file.unlink()
            logger.info(f"✅ {directory}: Read/write permissions OK")
        except Exception as e:
            logger.error(f"❌ {directory}: Permission error - {e}")
            return False
    
    logger.info("✅ All startup checks passed!")
    return True

def start_background_cleanup():
    """Start background cleanup thread"""
    def cleanup_worker():
        while True:
            try:
                # Clean up old cache files
                CacheManager.cleanup_old_cache()
                
                # Clean up temp files older than 1 hour
                cleanup_temp_files(max_age_hours=1)
                
                # Sleep for 30 minutes before next cleanup
                time.sleep(30 * 60)
                
            except Exception as e:
                logger.error(f"Background cleanup error: {e}")
                time.sleep(60)  # Wait 1 minute on error
    
    cleanup_thread = threading.Thread(target=cleanup_worker, daemon=True)
    cleanup_thread.start()
    logger.info("🧹 Background cleanup thread started")

if __name__ == '__main__':
    # Run startup checks
    if not run_startup_checks():
        logger.error("❌ Startup checks failed, exiting...")
        exit(1)
    
    # Start background cleanup
    start_background_cleanup()
    
    # Get port from environment or use default
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('FLASK_DEBUG', 'False').lower() == 'true'
    
    logger.info(f"🌐 Starting server on port {port} (debug={'on' if debug else 'off'})")
    
    # In production, use a proper WSGI server like Gunicorn
    # This is just for development/testing
    try:
        app.run(
            host='0.0.0.0',
            port=port,
            debug=debug,
            threaded=True,
            use_reloader=False  # Disable reloader in production
        )
    except Exception as e:
        logger.error(f"❌ Failed to start server: {e}")
        raise
